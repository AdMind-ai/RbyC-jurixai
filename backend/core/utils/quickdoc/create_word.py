from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
from django.conf import settings
from io import BytesIO

# -----------------------
# Função padrão com template
# -----------------------
def create_word_with_template(text, title):
    """
    Cria um Word usando um template (primeira página fixa)
    Adiciona título e corpo do texto seguindo lógica do PDF com template.
    """
    # Abrir template (se existir)
    template_path = Path(settings.STATIC_ROOT) / "quickdoc" / "template_word.docx"
    if template_path.exists():
        doc = Document(template_path)
    else:
        doc = Document()

    # -----------------------------
    # Espaço antes do título para não colar no header do template
    # -----------------------------
    header_to_title_space = 1  # parágrafos em branco antes do título

    # -----------------------------
    # Título centralizado
    # -----------------------------
    title_para = doc.add_paragraph("\n")
    title_run = title_para.add_run(title[:100])
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run.bold = True
    title_run.font.size = Pt(14)
    doc.add_paragraph()  # espaço após título

    # -----------------------------
    # Corpo do texto
    # -----------------------------
    max_lines_per_page = 31  # número aproximado de linhas por página
    
    lines = text.split("\n")
    line_count = 0
    first_paragraph_on_page = True

    for idx, line in enumerate(lines):
        if line.strip() == "":
            doc.add_paragraph()
            line_count += 1
        else:
            # Adiciona parágrafos em branco para simular altura do título
            if first_paragraph_on_page:
                for _ in range(header_to_title_space):
                    doc.add_paragraph("\n")
                first_paragraph_on_page = False

            # Adiciona parágrafo real
            para = doc.add_paragraph(line.strip())
            para.paragraph_format.line_spacing = 1.3
            para.paragraph_format.space_after = Pt(2)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.runs[0].font.size = Pt(11)
            line_count += 1

        # Quebra de página condicional (só se houver conteúdo restante)
        if line_count >= max_lines_per_page and idx < len(lines) - 1:
            doc.add_page_break()
            line_count = 0
            first_paragraph_on_page = True  # próxima página começa com espaço do título

    # -----------------------------
    # Salvar em BytesIO
    # -----------------------------
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer


# -----------------------
# Função verbale CDA
# -----------------------
def create_word_for_verbale_cda(text, title):
    """
    Cria um Word no estilo verbale CDA:
    - Header específico na primeira página
    - Título centralizado
    - Corpo do texto com quebra de página automática
    - Footer fixo
    """
    doc = Document()
    section = doc.sections[0]

    # Margens aproximadas ao PDF
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # ---------------------
    # HEADER (primeira página)
    # ---------------------
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    header_texts = [
        "Replica SIM S.p.A.",
        "Capitale sociale Euro 10.500.000 i.v.",
        "Codice fiscale e Partita IVA: 11064390963",
        "Sede legale in Milano - Corso Sempione, n. 2 – 20154 Milano"
    ]
    for t in header_texts:
        para = first_header.add_paragraph(t)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.runs[0]
        run.font.size = Pt(10)
        if "Replica SIM S.p.A." in t:
            run.bold = True

    doc.add_paragraph("\n")  

    # ---------------------
    # Título centralizado
    # ---------------------
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title[:100])
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run.bold = True
    title_run.font.size = Pt(14)
    doc.add_paragraph()  # espaço após título

    # ---------------------
    # Corpo do texto
    # ---------------------
    max_lines_per_page = 40
    line_count = 0
    for line in text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph()
            line_count += 1
        else:
            para = doc.add_paragraph(line.strip())
            para.paragraph_format.line_spacing = 1.3
            para.paragraph_format.space_after = Pt(2)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.runs[0].font.size = Pt(11)
            line_count += 1

        if line_count >= max_lines_per_page:
            doc.add_page_break()
            line_count = 0
        
    # ---------------------
    # Salvar em BytesIO
    # ---------------------
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer
