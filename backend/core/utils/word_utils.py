from docx import Document
from io import BytesIO
import base64
import re


def create_word_with_template(content: str, title: str, letterhead_base64: str = None):
    """Create a Word document from content and title.

    If `letterhead_base64` is provided and represents a valid .docx file (base64,
    optionally prefixed with `data:`), the function will open that template and
    append the generated content into it so the template's layout/header/footer
    are preserved.

    Returns a BytesIO buffer containing the generated .docx.
    """
    try:
        if letterhead_base64:
            lb = letterhead_base64
            if lb.startswith('data:'):
                lb = lb.split(',', 1)[1]
            template_bytes = base64.b64decode(lb)
            template_buf = BytesIO(template_bytes)
            try:
                doc = Document(template_buf)
            except Exception:
                # fallback to new document if template cannot be opened
                doc = Document()
        else:
            doc = Document()

        if title:
            try:
                doc.add_heading(title, level=1)
            except Exception:
                # Some templates may not include the standard 'Heading 1' style
                # (localized templates or minimal templates). Fall back to
                # inserting a bold paragraph instead of relying on styles.
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True

        # Split content into paragraphs on blank lines, then collapse any internal
        # line breaks within a paragraph into single spaces so long lines that
        # were wrapped do not become separate paragraphs in the .docx.
        content_str = str(content or '').strip()
        if content_str:
            paragraphs = re.split(r'\n\s*\n', content_str)
            for para in paragraphs:
                # Remove empty lines and trim each line, then join with space.
                lines = [ln.strip() for ln in para.splitlines() if ln.strip()]
                if not lines:
                    continue
                text = ' '.join(lines)
                doc.add_paragraph(text)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception as exc:
        raise RuntimeError(f'Word generation failed: {exc}')
